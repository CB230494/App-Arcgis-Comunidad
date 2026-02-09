# -*- coding: utf-8 -*-
# ==========================================================================================
# PARTE 1/5 — BASE + HELPERS + CATÁLOGO CANTÓN→DISTRITO + HEADER + ESTADO + SEED + SIDEBAR
#
# FIX CLAVE (para el error de XLSForm):
# - El problema viene porque `${acepta_participar}` se estaba usando en un `relevant`
#   de un grupo/pregunta ANTES de que la pregunta `acepta_participar` apareciera en la hoja survey.
# - En las siguientes partes, construiremos la página de Consentimiento en ORDEN CORRECTO:
#     begin_group -> note(consentimiento) -> pregunta(acepta_participar) -> end(relevant NO) -> end_group
#   y solo DESPUÉS, las demás páginas usarán `${acepta_participar}='si'`.
#
# IMPORTANTE:
# - Pegá las 5 partes EN ORDEN en un SOLO archivo `app.py` (una debajo de la otra).
# - No borres nada entre partes.
# ==========================================================================================

import re
import json
from io import BytesIO
from datetime import datetime
from typing import List, Dict

import streamlit as st
import pandas as pd

# ------------------------------------------------------------------------------------------
# Configuración de la app
# ------------------------------------------------------------------------------------------
st.set_page_config(page_title="Encuesta Comunidad → XLSForm (Survey123)", layout="wide")
st.title("🏘️ Encuesta Comunidad → XLSForm para ArcGIS Survey123")

st.markdown("""
Crea tu cuestionario y **exporta un XLSForm** listo para **ArcGIS Survey123**.

Incluye:
- Tipos: **text**, **integer**, **date**, **time**, **geopoint**, **select_one**, **select_multiple**.
- **Constructor completo** (agregar, editar, ordenar, borrar) con condicionales.
- **Listas en cascada** **Cantón→Distrito** (**catálogo manual por lotes**).
- **Páginas** con navegación **Siguiente/Anterior** (`settings.style = pages`).
- **Portada** con **logo** (`media::image`) e **introducción**.
- **Consentimiento informado** con finalización temprana si NO acepta.
- **Glosario por página** (opcional), editable y exportable.
""")

# ------------------------------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------------------------------
TIPOS = [
    "Texto (corto)",
    "Párrafo (texto largo)",
    "Número",
    "Selección única",
    "Selección múltiple",
    "Fecha",
    "Hora",
    "GPS (ubicación)"
]

def _rerun():
    """Compat: st.rerun() vs st.experimental_rerun()."""
    if hasattr(st, "rerun"):
        st.rerun()
    else:
        st.experimental_rerun()

def slugify_name(texto: str) -> str:
    """Convierte etiqueta humana a name seguro (solo a-z0-9_)."""
    if not texto:
        return "campo"
    t = texto.lower()
    t = re.sub(r"[áàäâ]", "a", t); t = re.sub(r"[éèëê]", "e", t)
    t = re.sub(r"[íìïî]", "i", t); t = re.sub(r"[óòöô]", "o", t)
    t = re.sub(r"[úùüû]", "u", t); t = re.sub(r"ñ", "n", t)
    t = re.sub(r"[^a-z0-9]+", "_", t).strip("_")
    return t or "campo"

def asegurar_nombre_unico(base: str, usados: set) -> str:
    """Evita colisiones de name (agrega _2, _3...)."""
    if base not in usados:
        return base
    i = 2
    while f"{base}_{i}" in usados:
        i += 1
    return f"{base}_{i}"

def map_tipo_to_xlsform(tipo_ui: str, name: str):
    """Mapea tipo UI -> XLSForm (type, appearance_default, list_name)."""
    if tipo_ui == "Texto (corto)":
        return ("text", None, None)
    if tipo_ui == "Párrafo (texto largo)":
        return ("text", "multiline", None)
    if tipo_ui == "Número":
        return ("integer", None, None)
    if tipo_ui == "Selección única":
        return (f"select_one list_{name}", None, f"list_{name}")
    if tipo_ui == "Selección múltiple":
        return (f"select_multiple list_{name}", None, f"list_{name}")
    if tipo_ui == "Fecha":
        return ("date", None, None)
    if tipo_ui == "Hora":
        return ("time", None, None)
    if tipo_ui == "GPS (ubicación)":
        return ("geopoint", None, None)
    return ("text", None, None)

def xlsform_or_expr(conds):
    if not conds:
        return None
    if len(conds) == 1:
        return conds[0]
    return "(" + " or ".join(conds) + ")"

def xlsform_not(expr):
    if not expr:
        return None
    return f"not({expr})"

def build_relevant_expr(rules_for_target: List[Dict]):
    """Convierte reglas {src,op,values[]} a XPath para relevant."""
    or_parts = []
    for r in rules_for_target:
        src = r["src"]
        op = r.get("op", "=")
        vals = r.get("values", [])
        if not vals:
            continue
        if op == "=":
            segs = [f"${{{src}}}='{v}'" for v in vals]
        elif op == "selected":
            segs = [f"selected(${{{src}}}, '{v}')" for v in vals]
        elif op == "!=":
            segs = [f"${{{src}}}!='{v}'" for v in vals]
        else:
            segs = [f"${{{src}}}='{v}'" for v in vals]
        or_parts.append(xlsform_or_expr(segs))
    return xlsform_or_expr(or_parts)

# ------------------------------------------------------------------------------------------
# Catálogo manual por lotes: Cantón → Distrito
# ------------------------------------------------------------------------------------------
if "choices_ext_rows" not in st.session_state:
    st.session_state.choices_ext_rows = []  # filas para hoja choices
if "choices_extra_cols" not in st.session_state:
    st.session_state.choices_extra_cols = set()

def _append_choice_unique(row: Dict):
    """Inserta fila en choices evitando duplicados por (list_name,name)."""
    key = (row.get("list_name"), row.get("name"))
    exists = any((r.get("list_name"), r.get("name")) == key for r in st.session_state.choices_ext_rows)
    if not exists:
        st.session_state.choices_ext_rows.append(row)

st.markdown("### 📚 Catálogo Cantón → Distrito (por lotes)")
with st.expander("Agrega un lote (un Cantón y uno o varios Distritos)", expanded=True):
    col_c1, col_c2 = st.columns(2)
    canton_txt = col_c1.text_input("Cantón (una vez)", value="")
    distritos_txt = col_c2.text_area("Distritos del cantón (uno por línea)", value="", height=130)

    col_b1, col_b2, col_b3 = st.columns([1, 1, 2])
    add_lote = col_b1.button("Agregar lote", type="primary", use_container_width=True)
    clear_all = col_b2.button("Limpiar catálogo", use_container_width=True)

    if clear_all:
        st.session_state.choices_ext_rows = []
        st.success("Catálogo limpiado.")

    if add_lote:
        c = canton_txt.strip()
        distritos = [d.strip() for d in distritos_txt.splitlines() if d.strip()]
        if not c or not distritos:
            st.error("Debes indicar Cantón y al menos un Distrito.")
        else:
            slug_c = slugify_name(c)

            # columnas extra usadas por filtros/placeholder
            st.session_state.choices_extra_cols.update({"canton_key", "any"})

            # Placeholders (una sola vez por lista)
            _append_choice_unique({"list_name": "list_canton", "name": "__pick_canton__", "label": "— escoja un cantón —"})
            _append_choice_unique({"list_name": "list_distrito", "name": "__pick_distrito__", "label": "— escoja un cantón —", "any": "1"})

            # Cantón
            _append_choice_unique({"list_name": "list_canton", "name": slug_c, "label": c})

            # Distritos
            usados_d = set()
            for d in distritos:
                slug_d = asegurar_nombre_unico(slugify_name(d), usados_d)
                usados_d.add(slug_d)
                _append_choice_unique({"list_name": "list_distrito", "name": slug_d, "label": d, "canton_key": slug_c})

            st.success(f"Lote agregado: {c} → {len(distritos)} distritos.")

# Vista previa de catálogo
if st.session_state.choices_ext_rows:
    st.dataframe(pd.DataFrame(st.session_state.choices_ext_rows),
                 use_container_width=True, hide_index=True, height=240)

# ------------------------------------------------------------------------------------------
# Cabecera: Logo + Delegación
# ------------------------------------------------------------------------------------------
DEFAULT_LOGO_PATH = "001.png"

col_logo, col_txt = st.columns([1, 3], vertical_alignment="center")
with col_logo:
    up_logo = st.file_uploader("Logo (PNG/JPG)", type=["png", "jpg", "jpeg"])
    if up_logo:
        st.image(up_logo, caption="Logo cargado", use_container_width=True)
        st.session_state["_logo_bytes"] = up_logo.getvalue()
        st.session_state["_logo_name"] = up_logo.name
    else:
        try:
            st.image(DEFAULT_LOGO_PATH, caption="Logo (001.png)", use_container_width=True)
            st.session_state["_logo_bytes"] = None
            st.session_state["_logo_name"] = "001.png"
        except Exception:
            st.warning("Sube un logo para incluirlo en el XLSForm.")
            st.session_state["_logo_bytes"] = None
            st.session_state["_logo_name"] = "logo.png"

with col_txt:
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    delegacion = st.text_input("Nombre del lugar / Delegación", value="San Carlos Oeste")
    logo_media_name = st.text_input(
        "Nombre de archivo para `media::image`",
        value=st.session_state.get("_logo_name", "001.png"),
        help="Debe coincidir con el archivo en `media/` de Survey123 Connect."
    )
    titulo_compuesto = (f"Encuesta comunidad – {delegacion.strip()}" if delegacion.strip() else "Encuesta comunidad")
    st.markdown(f"<h5 style='text-align:center;margin:4px 0'>📋 {titulo_compuesto}</h5>", unsafe_allow_html=True)

# ------------------------------------------------------------------------------------------
# Estado principal
# ------------------------------------------------------------------------------------------
if "preguntas" not in st.session_state:
    st.session_state.preguntas = []
if "reglas_visibilidad" not in st.session_state:
    st.session_state.reglas_visibilidad = []
if "reglas_finalizar" not in st.session_state:
    st.session_state.reglas_finalizar = []

# Glosario por página (opcional)
if "glosario_por_pagina" not in st.session_state:
    st.session_state.glosario_por_pagina = {
        "p1_intro": [],
        "p2_consent": [],
        "p3_demograficos": [],
        "p4_sentimiento": [],
        "p5_lugares": [],
        "p6_incidencia": [],
        "p7_riesgos": [],
        "p8_info_adicional": [],
    }

# ------------------------------------------------------------------------------------------
# Textos fijos: Intro + Consentimiento
# ------------------------------------------------------------------------------------------
INTRO_COMUNIDAD = (
    "Con el fin de hacer más segura nuestra comunidad, queremos concentrarnos en los problemas de "
    "seguridad más importantes. Por lo que debemos trabajar juntos, tanto con el gobierno local como "
    "con otras instituciones y la comunidad, para reducir los delitos y riesgos que afectan a la gente. "
    "Es importante recordar que la información que nos proporcionas es confidencial y solo se usará para "
    "mejorar la seguridad en nuestra área."
)

CONSENTIMIENTO_TEXTO = (
    "Consentimiento informado: Su participación es voluntaria. Puede omitir cualquier pregunta o "
    "finalizar la encuesta en cualquier momento. La información recopilada es confidencial y se utilizará "
    "exclusivamente con fines de diagnóstico y mejora de la seguridad comunitaria."
)

# ------------------------------------------------------------------------------------------
# Precarga (seed) — Incluye acepta_participar y SIN barrio
# ------------------------------------------------------------------------------------------
if "seed_cargado" not in st.session_state:
    v_mas_seguro = slugify_name("Más seguro")
    v_igual = slugify_name("Igual")
    v_menos_seg = slugify_name("Menos seguro")

    seed = [
        # ---------------- Consentimiento (P2) ----------------
        {
            "tipo_ui": "Selección única",
            "label": "¿Acepta participar en esta encuesta?",
            "name": "acepta_participar",
            "required": True,
            "opciones": ["Sí", "No"],
            "appearance": None,
            "choice_filter": None,
            "relevant": None
        },

        # ---------------- Datos demográficos (P3) ----------------
        {"tipo_ui": "Selección única", "label": "Cantón", "name": "canton", "required": True,
         "opciones": [], "appearance": None, "choice_filter": None, "relevant": None},
        {"tipo_ui": "Selección única", "label": "Distrito", "name": "distrito", "required": True,
         "opciones": [], "appearance": None, "choice_filter": "canton_key=${canton} or any='1'", "relevant": None},

        {"tipo_ui": "Número", "label": "Edad", "name": "edad", "required": True,
         "opciones": [], "appearance": None, "choice_filter": None, "relevant": None},
        {"tipo_ui": "Selección única", "label": "Género", "name": "genero", "required": True,
         "opciones": ["Masculino", "Femenino", "LGTBQ+"], "appearance": None, "choice_filter": None, "relevant": None},
        {"tipo_ui": "Selección única", "label": "Escolaridad", "name": "escolaridad", "required": True,
         "opciones": ["Ninguna", "Primaria", "Primaria incompleta", "Secundaria completa", "Secundaria incompleta",
                      "Universitaria", "Universitaria incompleta", "Técnico"],
         "appearance": None, "choice_filter": None, "relevant": None},
        {"tipo_ui": "Selección múltiple", "label": "¿Cuál es su relación con la zona?", "name": "relacion_zona", "required": True,
         "opciones": ["Vivo en la zona", "Trabajo en la zona", "Visito la zona"], "appearance": None, "choice_filter": None, "relevant": None},

        # ---------------- Sentimiento (P4) ----------------
        {"tipo_ui": "Selección única", "label": "¿Se siente seguro en su barrio?", "name": "se_siente_seguro", "required": True,
         "opciones": ["Si", "No"], "appearance": None, "choice_filter": None, "relevant": None},
        {"tipo_ui": "Párrafo (texto largo)", "label": "Indique por qué considera el barrio inseguro", "name": "motivo_inseguridad", "required": True,
         "opciones": [], "appearance": None, "choice_filter": None,
         "relevant": f"${{se_siente_seguro}}='{slugify_name('No')}'"},
        {"tipo_ui": "Selección única",
         "label": "¿Cómo se siente respecto a la seguridad en su barrio este año comparado con el anterior?",
         "name": "comparacion_anual", "required": True,
         "opciones": ["Más seguro", "Igual", "Menos seguro"], "appearance": None, "choice_filter": None, "relevant": None},
        {"tipo_ui": "Párrafo (texto largo)", "label": "Indique por qué.", "name": "motivo_comparacion", "required": True,
         "opciones": [], "appearance": None, "choice_filter": None,
         "relevant": xlsform_or_expr([
             f"${{comparacion_anual}}='{v_mas_seguro}'",
             f"${{comparacion_anual}}='{v_igual}'",
             f"${{comparacion_anual}}='{v_menos_seg}'"
         ])},
    ]

    st.session_state.preguntas = seed
    st.session_state.seed_cargado = True

# ------------------------------------------------------------------------------------------
# Sidebar: Metadatos + Exportar/Importar proyecto (JSON)  (incluye glosario + catálogo)
# ------------------------------------------------------------------------------------------
with st.sidebar:
    st.header("⚙️ Configuración")
    form_title = st.text_input(
        "Título del formulario",
        value=(f"Encuesta comunidad – {delegacion.strip()}" if delegacion.strip() else "Encuesta comunidad")
    )
    idioma = st.selectbox("Idioma por defecto (default_language)", options=["es", "en"], index=0)
    version_auto = datetime.now().strftime("%Y%m%d%H%M")
    version = st.text_input("Versión (settings.version)", value=version_auto)

    st.markdown("---")
    st.caption("💾 Exporta/Importa tu proyecto (JSON)")
    col_exp, col_imp = st.columns(2)

    if col_exp.button("Exportar proyecto (JSON)", use_container_width=True):
        proj = {
            "form_title": form_title,
            "idioma": idioma,
            "version": version,
            "preguntas": st.session_state.preguntas,
            "reglas_visibilidad": st.session_state.reglas_visibilidad,
            "reglas_finalizar": st.session_state.reglas_finalizar,
            "glosario_por_pagina": st.session_state.glosario_por_pagina,
            "choices_ext_rows": st.session_state.choices_ext_rows,
        }
        jbuf = BytesIO(json.dumps(proj, ensure_ascii=False, indent=2).encode("utf-8"))
        st.download_button(
            "Descargar JSON",
            data=jbuf,
            file_name="proyecto_encuesta.json",
            mime="application/json",
            use_container_width=True
        )

    up = col_imp.file_uploader("Importar JSON", type=["json"], label_visibility="collapsed")
    if up is not None:
        try:
            raw = up.read().decode("utf-8")
            data = json.loads(raw)
            st.session_state.preguntas = list(data.get("preguntas", []))
            st.session_state.reglas_visibilidad = list(data.get("reglas_visibilidad", []))
            st.session_state.reglas_finalizar = list(data.get("reglas_finalizar", []))
            st.session_state.glosario_por_pagina = dict(
                data.get("glosario_por_pagina", st.session_state.glosario_por_pagina)
            )
            st.session_state.choices_ext_rows = list(
                data.get("choices_ext_rows", st.session_state.choices_ext_rows)
            )
            _rerun()
        except Exception as e:
            st.error(f"No se pudo importar el JSON: {e}")

# ==========================================================================================
# FIN PARTE 1/5
# (Sigue PARTE 2/5: Constructor (Agregar/Editar/Ordenar/Borrar) + Glosario por página (UI))
# ==========================================================================================
# ==========================================================================================
# PARTE 2/5 — CONSTRUCTOR (AGREGAR/EDITAR/ORDENAR/BORRAR) + GLOSARIO POR PÁGINA (EDITOR)
# ==========================================================================================

# ------------------------------------------------------------------------------------------
# Constructor: Agregar nuevas preguntas
# ------------------------------------------------------------------------------------------
st.subheader("📝 Diseña tus preguntas")

with st.form("form_add_q", clear_on_submit=False):
    tipo_ui = st.selectbox("Tipo de pregunta", options=TIPOS)
    label = st.text_input("Etiqueta (texto exacto)")
    sugerido = slugify_name(label) if label else ""
    col_n1, col_n2, col_n3 = st.columns([2, 1, 1])
    name = col_n1.text_input("Nombre interno (XLSForm 'name')", value=sugerido)
    required = col_n2.checkbox("Requerida", value=False)
    appearance = col_n3.text_input("Appearance (opcional)", value="")

    opciones = []
    if tipo_ui in ("Selección única", "Selección múltiple"):
        st.markdown("**Opciones (una por línea)**")
        txt_opts = st.text_area("Opciones", height=120)
        if txt_opts.strip():
            opciones = [o.strip() for o in txt_opts.splitlines() if o.strip()]

    add = st.form_submit_button("➕ Agregar pregunta")

if add:
    if not label.strip():
        st.warning("Agrega una etiqueta.")
    else:
        base = slugify_name(name or label)
        usados = {q["name"] for q in st.session_state.preguntas}
        unico = asegurar_nombre_unico(base, usados)
        st.session_state.preguntas.append({
            "tipo_ui": tipo_ui,
            "label": label.strip(),
            "name": unico,
            "required": required,
            "opciones": opciones,
            "appearance": (appearance.strip() or None),
            "choice_filter": None,
            "relevant": None
        })
        st.success(f"Pregunta agregada: **{label}** (name: `{unico}`)")

# ------------------------------------------------------------------------------------------
# Lista / Ordenado / Edición (completa)
# ------------------------------------------------------------------------------------------
st.subheader("📚 Preguntas (ordénalas y edítalas)")
if not st.session_state.preguntas:
    st.info("Aún no has agregado preguntas.")
else:
    for idx, q in enumerate(st.session_state.preguntas):
        with st.container(border=True):
            c1, c2, c3, c4, c5 = st.columns([4, 2, 2, 2, 2])
            c1.markdown(f"**{idx+1}. {q['label']}**")

            meta = f"type: {q['tipo_ui']}  •  name: `{q['name']}`  •  requerida: {'sí' if q['required'] else 'no'}"
            if q.get("appearance"):
                meta += f"  •  appearance: `{q['appearance']}`"
            if q.get("choice_filter"):
                meta += f"  •  choice_filter: `{q['choice_filter']}`"
            if q.get("relevant"):
                meta += f"  •  relevant: `{q['relevant']}`"
            c1.caption(meta)

            if q["tipo_ui"] in ("Selección única", "Selección múltiple"):
                c1.caption("Opciones: " + ", ".join(q.get("opciones") or []))

            up = c2.button("⬆️ Subir", key=f"up_{idx}", use_container_width=True, disabled=(idx == 0))
            down = c3.button("⬇️ Bajar", key=f"down_{idx}", use_container_width=True,
                             disabled=(idx == len(st.session_state.preguntas) - 1))
            edit = c4.button("✏️ Editar", key=f"edit_{idx}", use_container_width=True)
            borrar = c5.button("🗑️ Eliminar", key=f"del_{idx}", use_container_width=True)

            if up:
                st.session_state.preguntas[idx - 1], st.session_state.preguntas[idx] = (
                    st.session_state.preguntas[idx], st.session_state.preguntas[idx - 1]
                )
                _rerun()

            if down:
                st.session_state.preguntas[idx + 1], st.session_state.preguntas[idx] = (
                    st.session_state.preguntas[idx], st.session_state.preguntas[idx + 1]
                )
                _rerun()

            if edit:
                st.markdown("**Editar esta pregunta**")
                ne_label = st.text_input("Etiqueta", value=q["label"], key=f"e_label_{idx}")
                ne_name = st.text_input("Nombre interno (name)", value=q["name"], key=f"e_name_{idx}")
                ne_required = st.checkbox("Requerida", value=q["required"], key=f"e_req_{idx}")
                ne_appearance = st.text_input("Appearance", value=q.get("appearance") or "", key=f"e_app_{idx}")
                ne_choice_filter = st.text_input("choice_filter (opcional)", value=q.get("choice_filter") or "", key=f"e_cf_{idx}")
                ne_relevant = st.text_input("relevant (opcional – se autogenera por reglas)", value=q.get("relevant") or "", key=f"e_rel_{idx}")

                ne_opciones = q.get("opciones") or []
                if q["tipo_ui"] in ("Selección única", "Selección múltiple"):
                    ne_opts_txt = st.text_area("Opciones (una por línea)", value="\n".join(ne_opciones), key=f"e_opts_{idx}")
                    ne_opciones = [o.strip() for o in ne_opts_txt.splitlines() if o.strip()]

                col_ok, col_cancel = st.columns(2)
                if col_ok.button("💾 Guardar cambios", key=f"e_save_{idx}", use_container_width=True):
                    new_base = slugify_name(ne_name or ne_label)
                    usados = {qq["name"] for j, qq in enumerate(st.session_state.preguntas) if j != idx}
                    ne_name_final = new_base if new_base not in usados else asegurar_nombre_unico(new_base, usados)

                    st.session_state.preguntas[idx]["label"] = ne_label.strip() or q["label"]
                    st.session_state.preguntas[idx]["name"] = ne_name_final
                    st.session_state.preguntas[idx]["required"] = ne_required
                    st.session_state.preguntas[idx]["appearance"] = ne_appearance.strip() or None
                    st.session_state.preguntas[idx]["choice_filter"] = ne_choice_filter.strip() or None
                    st.session_state.preguntas[idx]["relevant"] = ne_relevant.strip() or None
                    if q["tipo_ui"] in ("Selección única", "Selección múltiple"):
                        st.session_state.preguntas[idx]["opciones"] = ne_opciones

                    st.success("Cambios guardados.")
                    _rerun()

                if col_cancel.button("Cancelar", key=f"e_cancel_{idx}", use_container_width=True):
                    _rerun()

            if borrar:
                del st.session_state.preguntas[idx]
                st.warning("Pregunta eliminada.")
                _rerun()

# ------------------------------------------------------------------------------------------
# Glosario por página (editor)
# ------------------------------------------------------------------------------------------
st.markdown("---")
st.subheader("📘 Glosario por página (opcional)")

PAGES_UI = [
    ("p1_intro", "P1 — Introducción"),
    ("p2_consent", "P2 — Consentimiento"),
    ("p3_demograficos", "P3 — Datos demográficos"),
    ("p4_sentimiento", "P4 — Sentimiento de inseguridad"),
    ("p5_lugares", "P5 — Lugares del barrio"),
    ("p6_incidencia", "P6 — Incidencia relacionada a delitos"),
    ("p7_riesgos", "P7 — Riesgos sociales"),
    ("p8_info_adicional", "P8 — Información adicional"),
]

col_g1, col_g2 = st.columns([2, 3])
with col_g1:
    page_key = st.selectbox("Página", options=[k for k, _ in PAGES_UI], format_func=lambda k: dict(PAGES_UI)[k])
with col_g2:
    st.caption("Formato recomendado: una entrada por línea. Ejemplo: `Arrebato: robo rápido sin violencia…`")

current_list = st.session_state.glosario_por_pagina.get(page_key, [])
txt_current = "\n".join(current_list) if current_list else ""

new_txt = st.text_area("Entradas del glosario (una por línea)", value=txt_current, height=160)

col_ga, col_gb, col_gc = st.columns([1, 1, 2])
if col_ga.button("💾 Guardar glosario de esta página", type="primary", use_container_width=True):
    lines = [ln.strip() for ln in (new_txt or "").splitlines() if ln.strip()]
    # quitar duplicados conservando orden
    seen = set()
    cleaned = []
    for ln in lines:
        if ln not in seen:
            cleaned.append(ln)
            seen.add(ln)
    st.session_state.glosario_por_pagina[page_key] = cleaned
    st.success("Glosario guardado.")

if col_gb.button("🧹 Limpiar glosario de esta página", use_container_width=True):
    st.session_state.glosario_por_pagina[page_key] = []
    st.success("Glosario limpiado.")
    _rerun()

with st.expander("👀 Vista rápida (todas las páginas)", expanded=False):
    for k, title in PAGES_UI:
        items = st.session_state.glosario_por_pagina.get(k, [])
        st.markdown(f"**{title}**")
        if items:
            st.write("• " + "\n• ".join(items))
        else:
            st.caption("Sin entradas.")

# ==========================================================================================
# FIN PARTE 2/5
# (Sigue PARTE 3/5: Condicionales + Constructor XLSForm (con Consentimiento en orden correcto))
# ==========================================================================================
# ==========================================================================================
# PARTE 3/5 — CONDICIONALES + CONSTRUCCIÓN XLSFORM (CONSENTIMIENTO EN ORDEN CORRECTO)
#            + EXPORTAR XLSFORM (survey/choices/settings)
#
# AQUÍ VA EL FIX REAL DEL ERROR:
# - `acepta_participar` se escribe en la hoja survey ANTES de cualquier relevant que lo use.
# - Se agrega una regla de fin automático:
#     * un grupo "end_if_no" con relevant = ${acepta_participar}='no'
#     * adentro un NOTE final
#     * así Survey123 termina el flujo naturalmente (ya no muestra el resto).
# - Luego, todas las páginas P3..P8 quedan con relevant = ${acepta_participar}='si'
# ==========================================================================================

# ------------------------------------------------------------------------------------------
# Condicionales (mostrar / finalizar)
# ------------------------------------------------------------------------------------------
st.subheader("🔀 Condicionales (mostrar / finalizar)")

if not st.session_state.preguntas:
    st.info("Agrega preguntas para definir condicionales.")
else:
    with st.expander("👁️ Mostrar pregunta si se cumple condición", expanded=False):
        names = [q["name"] for q in st.session_state.preguntas]
        labels_by_name = {q["name"]: q["label"] for q in st.session_state.preguntas}

        target = st.selectbox(
            "Pregunta a mostrar (target)",
            options=names,
            format_func=lambda n: f"{n} — {labels_by_name[n]}",
            key="vis_target"
        )
        src = st.selectbox(
            "Depende de (source)",
            options=names,
            format_func=lambda n: f"{n} — {labels_by_name[n]}",
            key="vis_src"
        )
        op = st.selectbox("Operador", options=["=", "selected"], key="vis_op")

        src_q = next((q for q in st.session_state.preguntas if q["name"] == src), None)

        vals = []
        if src_q and src_q.get("opciones"):
            vals = st.multiselect("Valores (usa texto, internamente se usará slug)", options=src_q["opciones"], key="vis_vals")
            vals = [slugify_name(v) for v in vals]
        else:
            manual = st.text_input("Valor (si la pregunta no tiene opciones)", key="vis_manual")
            vals = [slugify_name(manual)] if manual.strip() else []

        if st.button("➕ Agregar regla de visibilidad", key="btn_add_vis"):
            if target == src:
                st.error("Target y Source no pueden ser la misma pregunta.")
            elif not vals:
                st.error("Indica al menos un valor.")
            else:
                st.session_state.reglas_visibilidad.append({"target": target, "src": src, "op": op, "values": vals})
                st.success("Regla agregada.")
                _rerun()

        if st.session_state.reglas_visibilidad:
            st.markdown("**Reglas de visibilidad actuales:**")
            for i, r in enumerate(st.session_state.reglas_visibilidad):
                st.write(f"- Mostrar **{r['target']}** si **{r['src']}** {r['op']} {r['values']}")
                if st.button(f"Eliminar regla #{i+1}", key=f"del_vis_{i}"):
                    del st.session_state.reglas_visibilidad[i]
                    _rerun()

    with st.expander("⏹️ Finalizar temprano si se cumple condición", expanded=False):
        names = [q["name"] for q in st.session_state.preguntas]
        labels_by_name = {q["name"]: q["label"] for q in st.session_state.preguntas}

        src2 = st.selectbox(
            "Condición basada en",
            options=names,
            format_func=lambda n: f"{n} — {labels_by_name[n]}",
            key="final_src"
        )
        op2 = st.selectbox("Operador", options=["=", "selected", "!="], key="final_op")

        src2_q = next((q for q in st.session_state.preguntas if q["name"] == src2), None)
        vals2 = []
        if src2_q and src2_q.get("opciones"):
            vals2 = st.multiselect("Valores (slug interno)", options=src2_q["opciones"], key="final_vals")
            vals2 = [slugify_name(v) for v in vals2]
        else:
            manual2 = st.text_input("Valor (si no hay opciones)", key="final_manual")
            vals2 = [slugify_name(manual2)] if manual2.strip() else []

        if st.button("➕ Agregar regla de finalización", key="btn_add_fin"):
            if not vals2:
                st.error("Indica al menos un valor.")
            else:
                idx_src = next((i for i, q in enumerate(st.session_state.preguntas) if q["name"] == src2), 0)
                st.session_state.reglas_finalizar.append({"src": src2, "op": op2, "values": vals2, "index_src": idx_src})
                st.success("Regla agregada.")
                _rerun()

        if st.session_state.reglas_finalizar:
            st.markdown("**Reglas de finalización actuales:**")
            for i, r in enumerate(st.session_state.reglas_finalizar):
                st.write(f"- Si **{r['src']}** {r['op']} {r['values']} ⇒ ocultar lo que sigue (efecto fin)")
                if st.button(f"Eliminar regla fin #{i+1}", key=f"del_fin_{i}"):
                    del st.session_state.reglas_finalizar[i]
                    _rerun()

# ------------------------------------------------------------------------------------------
# Construcción XLSForm
# ------------------------------------------------------------------------------------------
def _get_logo_media_name():
    return logo_media_name

def construir_xlsform(preguntas, form_title: str, idioma: str, version: str,
                      reglas_vis, reglas_fin, glosario_por_pagina: Dict[str, List[str]]):
    survey_rows = []
    choices_rows = []

    # 1) Index de preguntas por name
    by_name = {q["name"]: q for q in preguntas}

    # 2) Reglas visibilidad por target
    vis_by_target = {}
    for r in reglas_vis:
        vis_by_target.setdefault(r["target"], []).append(
            {"src": r["src"], "op": r.get("op", "="), "values": r.get("values", [])}
        )

    # 3) Finalización por índice (como ya lo tenías)
    fin_conds = []
    for r in reglas_fin:
        cond = build_relevant_expr([{"src": r["src"], "op": r.get("op", "="), "values": r.get("values", [])}])
        if cond:
            fin_conds.append((r["index_src"], cond))

    # 4) Helpers internos
    def _consent_yes_expr():
        # OJO: en choices se guarda slugify_name("Sí") => "si"
        return f"${{acepta_participar}}='{slugify_name('Sí')}'"

    def _consent_no_expr():
        return f"${{acepta_participar}}='{slugify_name('No')}'"

    def _add_note_glosario(page_key: str):
        items = glosario_por_pagina.get(page_key, []) or []
        if not items:
            return
        texto = "GLOSARIO:\n" + "\n".join([f"• {x}" for x in items])
        survey_rows.append({
            "type": "note",
            "name": f"glosario_{page_key}",
            "label": texto
        })

    def add_q(q, idx, force_relevant_prefix: str | None = None):
        """Agrega una pregunta a survey y (si aplica) sus opciones a choices."""
        x_type, default_app, list_name = map_tipo_to_xlsform(q["tipo_ui"], q["name"])

        # Relevant manual (si el user lo escribió en la pregunta)
        rel_manual = q.get("relevant") or None

        # Relevant por panel (reglas visibilidad)
        rel_panel = build_relevant_expr(vis_by_target.get(q["name"], []))

        # Relevant por finalizar temprano (reglas_fin)
        nots = [xlsform_not(cond) for idx_src, cond in fin_conds if idx_src < idx]
        rel_fin = "(" + " and ".join(nots) + ")" if nots else None

        # Prefijo obligatorio (ej: consentimiento SI para páginas posteriores)
        parts = [p for p in [force_relevant_prefix, rel_manual, rel_panel, rel_fin] if p]
        rel_final = parts[0] if (parts and len(parts) == 1) else (("(" + ") and (".join(parts) + ")") if parts else None)

        row = {"type": x_type, "name": q["name"], "label": q["label"]}
        if q.get("required"):
            row["required"] = "yes"

        app = q.get("appearance") or default_app
        if app:
            row["appearance"] = app

        if q.get("choice_filter"):
            row["choice_filter"] = q["choice_filter"]

        if rel_final:
            row["relevant"] = rel_final

        # Constraints placeholders para cantón/distrito
        if q["name"] == "canton":
            row["constraint"] = ". != '__pick_canton__'"
            row["constraint_message"] = "Seleccione un cantón válido."
        if q["name"] == "distrito":
            row["constraint"] = ". != '__pick_distrito__'"
            row["constraint_message"] = "Seleccione un distrito válido."

        survey_rows.append(row)

        # Choices: NO generar opciones para cantón/distrito (usan catálogo manual)
        if list_name and q["name"] not in {"canton", "distrito"}:
            usados = set()
            for opt_label in (q.get("opciones") or []):
                base = slugify_name(opt_label)
                opt_name = asegurar_nombre_unico(base, usados)
                usados.add(opt_name)
                choices_rows.append({"list_name": list_name, "name": opt_name, "label": str(opt_label)})

    # --------------------------------------------------------------------------------------
    # PÁGINA 1: INTRO
    # --------------------------------------------------------------------------------------
    survey_rows += [
        {"type": "begin_group", "name": "p1_intro", "label": "Introducción", "appearance": "field-list"},
        {"type": "note", "name": "intro_logo", "label": form_title, "media::image": _get_logo_media_name()},
        {"type": "note", "name": "intro_texto", "label": INTRO_COMUNIDAD},
    ]
    _add_note_glosario("p1_intro")
    survey_rows += [
        {"type": "end_group", "name": "p1_intro_end"}
    ]

    # --------------------------------------------------------------------------------------
    # PÁGINA 2: CONSENTIMIENTO (ORDEN CORRECTO)
    # --------------------------------------------------------------------------------------
    survey_rows.append({"type": "begin_group", "name": "p2_consent", "label": "Consentimiento informado", "appearance": "field-list"})
    survey_rows.append({"type": "note", "name": "consent_texto", "label": CONSENTIMIENTO_TEXTO})
    _add_note_glosario("p2_consent")

    # Agregar la pregunta acepta_participar (SIEMPRE antes de usarla en relevant)
    if "acepta_participar" in by_name:
        idx_q = next((i for i, qq in enumerate(preguntas) if qq["name"] == "acepta_participar"), 0)
        add_q(by_name["acepta_participar"], idx_q, force_relevant_prefix=None)
    else:
        # Si alguien la borró, igual no rompemos. (pero Survey123 no podrá filtrar)
        survey_rows.append({"type": "note", "name": "warn_no_consent", "label": "⚠️ Falta la pregunta 'acepta_participar'."})

    survey_rows.append({"type": "end_group", "name": "p2_consent_end"})

    # Grupo de cierre si NO acepta (fin natural del flujo)
    survey_rows.append({"type": "begin_group", "name": "end_if_no", "label": "Fin", "appearance": "field-list", "relevant": _consent_no_expr()})
    survey_rows.append({"type": "note", "name": "fin_no", "label": "Gracias. Como no aceptó participar, la encuesta finaliza aquí."})
    survey_rows.append({"type": "end_group", "name": "end_if_no_end"})

    # --------------------------------------------------------------------------------------
    # Define sets por página (en esta versión, P3..P8)
    # --------------------------------------------------------------------------------------
    # P3: demográficos (sin barrio)
    p3 = {"canton", "distrito", "edad", "genero", "escolaridad", "relacion_zona"}

    # P4: sentimiento
    p4 = {"se_siente_seguro", "motivo_inseguridad", "comparacion_anual", "motivo_comparacion"}

    # Estas páginas no vienen completas en el seed de Parte 1 (pero tu código real sí las tendrá).
    # Igual dejamos la estructura; si no existen preguntas con esos name, simplemente no se agregan.
    p5 = {
        "lugar_entretenimiento", "espacios_recreativos", "lugar_residencia", "paradas_estaciones",
        "puentes_peatonales", "transporte_publico", "zona_bancaria", "zona_comercio",
        "zonas_residenciales", "lugares_turisticos", "zona_mas_insegura", "porque_insegura"
    }
    p6 = {
        "incidencia_delitos", "venta_drogas", "delitos_vida", "delitos_sexuales", "asaltos", "estafas",
        "robo_fuerza", "abandono_personas", "explotacion_infantil", "delitos_ambientales", "trata_personas",
        "vi", "vi_victima_ultimo_anno", "vi_tipos", "vi_fp_abordaje", "vi_fp_eval"
    }
    p7 = {"riesgos_sociales", "falta_inversion_social", "consumo_drogas", "infra_vial", "bunker"}
    p8 = {
        "info_grupo_delito", "desc_info_grupo", "victimizacion_12m",
        "delito_victima_si", "modo_operar_si", "horario_hecho_si",
        "delito_victima_no", "motivo_no_denuncia", "modo_operar_no", "horario_hecho_no",
        "fp_calificacion", "fp_24m", "conoce_policias", "conversa_policias",
        "sugerencia_fp", "sugerencia_muni", "otra_info", "contacto_voluntario"
    }

    def add_page(group_name: str, page_label: str, names_set: set, glos_key: str):
        survey_rows.append({
            "type": "begin_group",
            "name": group_name,
            "label": page_label,
            "appearance": "field-list",
            # IMPORTANTE: todas las páginas posteriores exigen consentimiento = Sí
            "relevant": _consent_yes_expr()
        })
        # glosario dentro de la página (arriba)
        _add_note_glosario(glos_key)

        for i, q in enumerate(preguntas):
            if q["name"] in names_set and q["name"] != "acepta_participar":
                add_q(q, i, force_relevant_prefix=None)  # el grupo ya tiene relevant por consentimiento
        survey_rows.append({"type": "end_group", "name": f"{group_name}_end"})

    add_page("p3_demograficos", "Datos demográficos", p3, "p3_demograficos")
    add_page("p4_sentimiento", "Sentimiento de inseguridad en el barrio", p4, "p4_sentimiento")
    add_page("p5_lugares", "Indique cómo se siente en los siguientes lugares de su barrio", p5, "p5_lugares")
    add_page("p6_incidencia", "Incidencia relacionada a delitos", p6, "p6_incidencia")
    add_page("p7_riesgos", "Riesgos Sociales", p7, "p7_riesgos")
    add_page("p8_info_adicional", "Información adicional", p8, "p8_info_adicional")

    # --------------------------------------------------------------------------------------
    # Choices del catálogo manual (con unicidad por list+name)
    # --------------------------------------------------------------------------------------
    for r in st.session_state.choices_ext_rows:
        choices_rows.append(dict(r))

    # --------------------------------------------------------------------------------------
    # DataFrames (survey/choices/settings)
    # --------------------------------------------------------------------------------------
    survey_cols_all = set().union(*[r.keys() for r in survey_rows]) if survey_rows else set()
    survey_cols = [c for c in [
        "type", "name", "label", "required", "appearance", "choice_filter", "relevant",
        "constraint", "constraint_message", "media::image"
    ] if c in survey_cols_all]
    for k in sorted(survey_cols_all):
        if k not in survey_cols:
            survey_cols.append(k)
    df_survey = pd.DataFrame(survey_rows, columns=survey_cols)

    choices_cols_all = set()
    for r in choices_rows:
        choices_cols_all.update(r.keys())
    base_choice_cols = ["list_name", "name", "label"]
    for extra in sorted(choices_cols_all):
        if extra not in base_choice_cols:
            base_choice_cols.append(extra)
    df_choices = pd.DataFrame(choices_rows, columns=base_choice_cols) if choices_rows else pd.DataFrame(columns=base_choice_cols)

    df_settings = pd.DataFrame([{
        "form_title": form_title,
        "version": version,
        "default_language": idioma,
        "style": "pages"
    }], columns=["form_title", "version", "default_language", "style"])

    return df_survey, df_choices, df_settings

def descargar_excel_xlsform(df_survey, df_choices, df_settings, nombre_archivo: str):
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df_survey.to_excel(writer, sheet_name="survey", index=False)
        df_choices.to_excel(writer, sheet_name="choices", index=False)
        df_settings.to_excel(writer, sheet_name="settings", index=False)

        wb = writer.book
        fmt_hdr = wb.add_format({"bold": True, "align": "left"})
        for sheet, df in (("survey", df_survey), ("choices", df_choices), ("settings", df_settings)):
            ws = writer.sheets[sheet]
            ws.freeze_panes(1, 0)
            ws.set_row(0, None, fmt_hdr)
            for col_idx, col_name in enumerate(list(df.columns)):
                ws.set_column(col_idx, col_idx, max(14, min(42, len(str(col_name)) + 8)))
    buffer.seek(0)

    st.download_button(
        label=f"📥 Descargar XLSForm ({nombre_archivo})",
        data=buffer,
        file_name=nombre_archivo,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )

# ------------------------------------------------------------------------------------------
# Exportar / Vista previa XLSForm
# ------------------------------------------------------------------------------------------
st.markdown("---")
st.subheader("📦 Generar XLSForm (Excel) para Survey123")
st.caption("""
Incluye:
- **survey** con `type,name,label,required,appearance,choice_filter,relevant,constraint,media::image`,
- **choices** (con `canton_key` y `any` para placeholders),
- **settings** con título, versión, idioma y **style = pages**.
""")

if st.button("🧮 Construir XLSForm", use_container_width=True, disabled=not st.session_state.preguntas):
    try:
        names = [q["name"] for q in st.session_state.preguntas]
        if len(names) != len(set(names)):
            st.error("Hay 'name' duplicados. Edita las preguntas para que cada 'name' sea único.")
        else:
            df_survey, df_choices, df_settings = construir_xlsform(
                st.session_state.preguntas,
                form_title=(f"Encuesta comunidad – {delegacion.strip()}" if delegacion.strip() else "Encuesta comunidad"),
                idioma=idioma,
                version=(version.strip() or datetime.now().strftime("%Y%m%d%H%M")),
                reglas_vis=st.session_state.reglas_visibilidad,
                reglas_fin=st.session_state.reglas_finalizar,
                glosario_por_pagina=st.session_state.glosario_por_pagina
            )
            st.success("XLSForm construido. Vista previa:")
            c1, c2, c3 = st.columns(3)
            c1.markdown("**Hoja: survey**");   c1.dataframe(df_survey, use_container_width=True, hide_index=True)
            c2.markdown("**Hoja: choices**");  c2.dataframe(df_choices, use_container_width=True, hide_index=True)
            c3.markdown("**Hoja: settings**"); c3.dataframe(df_settings, use_container_width=True, hide_index=True)

            nombre_archivo = slugify_name(form_title) + "_xlsform.xlsx"
            descargar_excel_xlsform(df_survey, df_choices, df_settings, nombre_archivo)

            if st.session_state.get("_logo_bytes"):
                st.download_button(
                    "📥 Descargar logo para carpeta media",
                    data=st.session_state["_logo_bytes"],
                    file_name=logo_media_name,
                    mime="image/png",
                    use_container_width=True
                )

            st.info("Publica en Survey123 Connect: crea encuesta desde archivo, copia el logo a `media/` y publica.")
    except Exception as e:
        st.error(f"Ocurrió un error al generar el XLSForm: {e}")

# ==========================================================================================
# FIN PARTE 3/5
# (Sigue PARTE 4/5: Exportar Word y PDF (manteniendo tu formato) + glosario por páginas)
# ==========================================================================================
# ==========================================================================================
# PARTE 4/5 — EXPORTAR WORD + PDF EDITABLE (P1 + P4..P8; SIN P3 si así lo querés)
#            + (mantiene tu estilo) + integra glosario por página en los documentos
#
# NOTA:
# - Mantenemos el mismo patrón que tu versión “perfecta”.
# - Imprimimos INTRO (P1) y las páginas que definimos (P4..P8) como en tu lógica original.
# - Puedes cambiar qué páginas se imprimen modificando ALL_BY_PAGE_PRINT abajo.
# ==========================================================================================

# ------------------------------------------------------------------------------------------
# Exportar Word y PDF
# ------------------------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.lib.colors import HexColor, black
except Exception:
    canvas = None

def _build_cond_text(qname: str, reglas_vis: List[Dict]) -> str:
    rels = [r for r in reglas_vis if r.get("target") == qname]
    if not rels:
        return ""
    parts = []
    for r in rels:
        op = r.get("op", "=")
        vals = r.get("values", [])
        vtxt = ", ".join(vals) if vals else ""
        parts.append(f"{r['src']} {op} [{vtxt}]")
    return "Condición: se muestra si " + " OR ".join(parts)

def _get_logo_bytes_fallback() -> bytes | None:
    if st.session_state.get("_logo_bytes"):
        return st.session_state["_logo_bytes"]
    try:
        with open("001.png", "rb") as f:
            return f.read()
    except Exception:
        return None

def _wrap_text_lines(text: str, font_name: str, font_size: float, max_width: float) -> List[str]:
    if not text:
        return []
    from reportlab.pdfbase.pdfmetrics import stringWidth
    words = text.split()
    lines, current = [], ""
    for w in words:
        test = (current + " " + w).strip()
        if stringWidth(test, font_name, font_size) <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            if stringWidth(w, font_name, font_size) > max_width:
                chunk = ""
                for ch in w:
                    if stringWidth(chunk + ch, font_name, font_size) <= max_width:
                        chunk += ch
                    else:
                        if chunk:
                            lines.append(chunk)
                        chunk = ch
                current = chunk
            else:
                current = w
    if current:
        lines.append(current)
    return lines

def _is_yes_no_options(opts: List[str]) -> bool:
    if not opts:
        return False
    norm = {slugify_name(x) for x in opts if x and str(x).strip()}
    yes_variants = {"si", "sí", "yes"}
    no_variants = {"no"}
    return norm.issubset(yes_variants | no_variants) and any(y in norm for y in yes_variants) and any(n in norm for n in no_variants)

def _should_show_options(q: Dict) -> bool:
    if q.get("tipo_ui") not in ("Selección única", "Selección múltiple"):
        return False
    opts = q.get("opciones") or []
    return bool(opts) and not _is_yes_no_options(opts)

def _glosario_text(page_key: str) -> str:
    items = st.session_state.glosario_por_pagina.get(page_key, []) or []
    if not items:
        return ""
    return "GLOSARIO:\n" + "\n".join([f"• {x}" for x in items])

# ------------------------------------------------------------------------------------------
# Definición de páginas (por name) para impresión
# (Ajusta libremente si querés imprimir otras páginas)
# ------------------------------------------------------------------------------------------
P4_NAMES = {"se_siente_seguro", "motivo_inseguridad", "comparacion_anual", "motivo_comparacion"}

P5_NAMES = {
    "lugar_entretenimiento", "espacios_recreativos", "lugar_residencia", "paradas_estaciones",
    "puentes_peatonales", "transporte_publico", "zona_bancaria", "zona_comercio",
    "zonas_residenciales", "lugares_turisticos", "zona_mas_insegura", "porque_insegura"
}

P6_NAMES = {
    "incidencia_delitos", "venta_drogas", "delitos_vida", "delitos_sexuales", "asaltos", "estafas",
    "robo_fuerza", "abandono_personas", "explotacion_infantil", "delitos_ambientales", "trata_personas",
    "vi", "vi_victima_ultimo_anno", "vi_tipos", "vi_fp_abordaje", "vi_fp_eval"
}

P7_NAMES = {"riesgos_sociales", "falta_inversion_social", "consumo_drogas", "infra_vial", "bunker"}

P8_NAMES = {
    "info_grupo_delito", "desc_info_grupo", "victimizacion_12m",
    "delito_victima_si", "modo_operar_si", "horario_hecho_si",
    "delito_victima_no", "motivo_no_denuncia", "modo_operar_no", "horario_hecho_no",
    "fp_calificacion", "fp_24m", "conoce_policias", "conversa_policias",
    "sugerencia_fp", "sugerencia_muni", "otra_info", "contacto_voluntario"
}

ALL_BY_PAGE_PRINT = [
    ("Introducción (P1)", "__INTRO__", "p1_intro"),
    ("Consentimiento (P2)", "__CONSENT__", "p2_consent"),
    ("Sentimiento de inseguridad (P4)", P4_NAMES, "p4_sentimiento"),
    ("Lugares del barrio (P5)", P5_NAMES, "p5_lugares"),
    ("Incidencia delictiva (P6)", P6_NAMES, "p6_incidencia"),
    ("Riesgos Sociales (P7)", P7_NAMES, "p7_riesgos"),
    ("Información adicional (P8)", P8_NAMES, "p8_info_adicional"),
]

# ------------------------------------------------------------------------------------------
# Helpers docx (cajas de observación)
# ------------------------------------------------------------------------------------------
def _set_cell_shading(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.replace('#', '').upper())

def _set_cell_borders(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '8')
        tag.set(qn('w:color'), color_hex.replace('#', '').upper())
        borders.append(tag)

def _add_observation_box(doc: Document, fill_hex: str, border_hex: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    _set_cell_shading(cell, fill_hex)
    _set_cell_borders(cell, border_hex)
    row = tbl.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = Inches(1.1)
    cell.paragraphs[0].add_run("")

# ------------------------------------------------------------------------------------------
# Export DOCX
# ------------------------------------------------------------------------------------------
def export_docx_form(preguntas: List[Dict], form_title: str, intro: str, consentimiento: str, reglas_vis: List[Dict]):
    if Document is None:
        st.error("Falta dependencia: instala `python-docx` para generar Word.")
        return

    fills = ["#E6F4EA", "#E7F0FE", "#FDECEA"]
    borders = ["#1E8E3E", "#1A73E8", "#D93025"]
    BLACK = RGBColor(0, 0, 0)

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(form_title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLACK
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    logo_b = _get_logo_bytes_fallback()
    if logo_b:
        try:
            img_buf = BytesIO(logo_b)
            doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_picture(img_buf, width=Inches(2.8))
        except Exception:
            pass

    # Intro
    intro_p = doc.add_paragraph(intro)
    if intro_p.runs:
        intro_p.runs[0].font.size = Pt(12)
        intro_p.runs[0].font.color.rgb = BLACK

    gtxt = _glosario_text("p1_intro")
    if gtxt:
        gp = doc.add_paragraph(gtxt)
        if gp.runs:
            gp.runs[0].font.size = Pt(10)
            gp.runs[0].font.color.rgb = BLACK

    # Consentimiento
    doc.add_paragraph("")
    sec = doc.add_paragraph("Consentimiento informado")
    rs = sec.runs[0]
    rs.bold = True
    rs.font.size = Pt(14)
    rs.font.color.rgb = BLACK

    cp = doc.add_paragraph(consentimiento)
    if cp.runs:
        cp.runs[0].font.size = Pt(12)
        cp.runs[0].font.color.rgb = BLACK

    gtxt = _glosario_text("p2_consent")
    if gtxt:
        gp = doc.add_paragraph(gtxt)
        if gp.runs:
            gp.runs[0].font.size = Pt(10)
            gp.runs[0].font.color.rgb = BLACK

    # Preguntas por secciones
    i = 1
    color_idx = 0
    for section_title, names, page_key in ALL_BY_PAGE_PRINT:
        if names in ("__INTRO__", "__CONSENT__"):
            continue

        doc.add_paragraph("")
        sec = doc.add_paragraph(section_title)
        rs = sec.runs[0]
        rs.bold = True
        rs.font.size = Pt(14)
        rs.font.color.rgb = BLACK

        gtxt = _glosario_text(page_key)
        if gtxt:
            gp = doc.add_paragraph(gtxt)
            if gp.runs:
                gp.runs[0].font.size = Pt(10)
                gp.runs[0].font.color.rgb = BLACK

        for q in preguntas:
            if q.get("name") not in names:
                continue

            doc.add_paragraph("")
            h = doc.add_paragraph(f"{i}. {q['label']}")
            if h.runs:
                r = h.runs[0]
                r.font.size = Pt(11)
                r.font.color.rgb = BLACK

            cond_txt = _build_cond_text(q["name"], reglas_vis)
            if cond_txt:
                cpara = doc.add_paragraph(cond_txt)
                if cpara.runs:
                    rc = cpara.runs[0]
                    rc.italic = True
                    rc.font.size = Pt(9)
                    rc.font.color.rgb = BLACK

            if _should_show_options(q):
                opts_str = ", ".join([str(x) for x in q.get("opciones") if str(x).strip()])
                opara = doc.add_paragraph(f"Opciones: {opts_str}")
                if opara.runs:
                    ro = opara.runs[0]
                    ro.font.size = Pt(10)
                    ro.font.color.rgb = BLACK

            fill = fills[color_idx % len(fills)]
            border = borders[color_idx % len(borders)]
            color_idx += 1
            _add_observation_box(doc, fill, border)

            help_p = doc.add_paragraph("Agregue sus observaciones sobre la pregunta.")
            if help_p.runs:
                rh = help_p.runs[0]
                rh.italic = True
                rh.font.size = Pt(9)
                rh.font.color.rgb = BLACK

            i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    st.download_button(
        "📄 Descargar Word del formulario",
        data=buf,
        file_name=slugify_name(form_title) + "_formulario.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

# ------------------------------------------------------------------------------------------
# Export PDF editable
# ------------------------------------------------------------------------------------------
def export_pdf_editable_form(preguntas: List[Dict], form_title: str, intro: str, consentimiento: str, reglas_vis: List[Dict]):
    if canvas is None:
        st.error("Falta dependencia: instala `reportlab` para generar PDF.")
        return

    from reportlab.pdfgen import canvas as _canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.utils import ImageReader
    from reportlab.lib.colors import HexColor, black

    PAGE_W, PAGE_H = A4
    margin = 2 * cm
    max_text_w = PAGE_W - 2 * margin

    title_font, title_size = "Helvetica-Bold", 24
    intro_font, intro_size = "Helvetica", 12
    intro_line_h = 18
    sec_font, sec_size = "Helvetica-Bold", 14
    label_font, label_size = "Helvetica", 11
    cond_font, cond_size = "Helvetica-Oblique", 9
    helper_font, helper_size = "Helvetica-Oblique", 9
    opts_font, opts_size = "Helvetica", 10

    fills = [HexColor("#E6F4EA"), HexColor("#E7F0FE"), HexColor("#FDECEA")]
    borders = [HexColor("#1E8E3E"), HexColor("#1A73E8"), HexColor("#D93025")]

    field_h = 80
    line_h = 14
    y = PAGE_H - margin

    c = _canvas.Canvas(BytesIO(), pagesize=A4)
    buf = c._filename
    c = _canvas.Canvas(buf, pagesize=A4)
    c.setTitle(form_title)

    # Logo
    logo_b = _get_logo_bytes_fallback()
    if logo_b:
        try:
            img = ImageReader(BytesIO(logo_b))
            logo_w, logo_h = 160, 115
            c.drawImage(img, (PAGE_W - logo_w) / 2, y - logo_h, width=logo_w, height=logo_h,
                        preserveAspectRatio=True, mask='auto')
            y -= (logo_h + 24)
        except Exception:
            pass

    # Título
    c.setFillColor(black)
    c.setFont(title_font, title_size)
    c.drawCentredString(PAGE_W / 2, y, form_title)
    y -= 26

    # Intro
    c.setFont(intro_font, intro_size)
    for line in _wrap_text_lines(intro, intro_font, intro_size, max_text_w):
        if y < margin + 80:
            c.showPage()
            y = PAGE_H - margin
            c.setFillColor(black)
            c.setFont(intro_font, intro_size)
        c.drawString(margin, y, line)
        y -= intro_line_h

    gtxt = _glosario_text("p1_intro")
    if gtxt:
        y -= 6
        for line in _wrap_text_lines(gtxt, "Helvetica", 10, max_text_w):
            if y < margin + 80:
                c.showPage()
                y = PAGE_H - margin
                c.setFillColor(black)
            c.setFont("Helvetica", 10)
            c.drawString(margin, y, line)
            y -= 14

    # Consentimiento
    y -= 8
    c.setFont(sec_font, sec_size)
    c.drawString(margin, y, "Consentimiento informado")
    y -= 18
    c.setFont(intro_font, intro_size)
    for line in _wrap_text_lines(consentimiento, intro_font, intro_size, max_text_w):
        if y < margin + 80:
            c.showPage()
            y = PAGE_H - margin
            c.setFillColor(black)
        c.drawString(margin, y, line)
        y -= intro_line_h

    gtxt = _glosario_text("p2_consent")
    if gtxt:
        y -= 6
        for line in _wrap_text_lines(gtxt, "Helvetica", 10, max_text_w):
            if y < margin + 80:
                c.showPage()
                y = PAGE_H - margin
                c.setFillColor(black)
            c.setFont("Helvetica", 10)
            c.drawString(margin, y, line)
            y -= 14

    c.showPage()
    y = PAGE_H - margin
    c.setFillColor(black)

    # Preguntas
    i = 1
    color_idx = 0
    for section_title, names, page_key in ALL_BY_PAGE_PRINT:
        if names in ("__INTRO__", "__CONSENT__"):
            continue

        c.setFont(sec_font, sec_size)
        c.drawString(margin, y, section_title)
        y -= (line_h + 6)

        gtxt = _glosario_text(page_key)
        if gtxt:
            c.setFont("Helvetica", 10)
            for line in _wrap_text_lines(gtxt, "Helvetica", 10, max_text_w):
                if y < margin + 120:
                    c.showPage()
                    y = PAGE_H - margin
                    c.setFillColor(black)
                    c.setFont(sec_font, sec_size)
                    c.drawString(margin, y, section_title)
                    y -= (line_h + 6)
                    c.setFont("Helvetica", 10)
                c.drawString(margin, y, line)
                y -= 14
            y -= 6

        c.setFont(label_font, label_size)

        for q in st.session_state.preguntas:
            if q.get("name") not in names:
                continue

            label_lines = _wrap_text_lines(f"{i}. {q['label']}", label_font, label_size, max_text_w)
            needed = line_h * len(label_lines) + field_h + 26

            cond_txt = _build_cond_text(q["name"], reglas_vis)
            cond_lines = _wrap_text_lines(cond_txt, cond_font, cond_size, max_text_w) if cond_txt else []
            needed += line_h * len(cond_lines)

            opts_lines = []
            if _should_show_options(q):
                opts_str = ", ".join([str(x) for x in q.get("opciones") if str(x).strip()])
                opts_lines = _wrap_text_lines(f"Opciones: {opts_str}", opts_font, opts_size, max_text_w)
                needed += line_h * len(opts_lines)

            if y - needed < margin:
                c.showPage()
                y = PAGE_H - margin
                c.setFillColor(black)
                c.setFont(sec_font, sec_size)
                c.drawString(margin, y, section_title)
                y -= (line_h + 6)
                c.setFont(label_font, label_size)

            for line in label_lines:
                c.drawString(margin, y, line)
                y -= line_h

            if cond_lines:
                c.setFont(cond_font, cond_size)
                for cl in cond_lines:
                    c.drawString(margin, y, cl)
                    y -= line_h
                c.setFont(label_font, label_size)

            if opts_lines:
                c.setFont(opts_font, opts_size)
                for ol in opts_lines:
                    c.drawString(margin, y, ol)
                    y -= line_h
                c.setFont(label_font, label_size)

            fill_color = fills[color_idx % len(fills)]
            border_color = borders[color_idx % len(borders)]
            color_idx += 1

            c.setFillColor(fill_color)
            c.setStrokeColor(border_color)
            c.rect(margin, y - field_h, max_text_w, field_h, fill=1, stroke=1)
            c.setFillColor(black)

            c.acroForm.textfield(
                name=f"campo_obs_{i}",
                tooltip=f"Observaciones para: {q['name']}",
                x=margin, y=y - field_h,
                width=max_text_w, height=field_h,
                borderWidth=1, borderStyle='solid', forceBorder=True,
                fieldFlags=4096, value=""
            )
            c.setFont(helper_font, helper_size)
            c.drawString(margin, y - field_h - 10, "Agregue sus observaciones sobre la pregunta.")
            c.setFont(label_font, label_size)

            y -= (field_h + 26)
            i += 1

        if y < margin + 120:
            c.showPage()
            y = PAGE_H - margin
            c.setFillColor(black)

    c.showPage()
    c.save()

    pdf_buf = c._filename
    data = pdf_buf.getvalue() if hasattr(pdf_buf, "getvalue") else pdf_buf
    st.download_button(
        "🧾 Descargar PDF editable del formulario",
        data=data,
        file_name=slugify_name(form_title) + "_formulario_editable.pdf",
        mime="application/pdf",
        use_container_width=True
    )

# ------------------------------------------------------------------------------------------
# Botones
# ------------------------------------------------------------------------------------------
st.markdown("### 📝 Exportar formulario en **Word** y **PDF editable**")
col_w, col_p = st.columns(2)

if col_w.button("Generar Word (DOCX)"):
    export_docx_form(
        st.session_state.preguntas,
        form_title=(f"Encuesta comunidad – {delegacion.strip()}" if delegacion.strip() else "Encuesta comunidad"),
        intro=INTRO_COMUNIDAD,
        consentimiento=CONSENTIMIENTO_TEXTO,
        reglas_vis=st.session_state.reglas_visibilidad
    )

if col_p.button("Generar PDF editable"):
    export_pdf_editable_form(
        st.session_state.preguntas,
        form_title=(f"Encuesta comunidad – {delegacion.strip()}" if delegacion.strip() else "Encuesta comunidad"),
        intro=INTRO_COMUNIDAD,
        consentimiento=CONSENTIMIENTO_TEXTO,
        reglas_vis=st.session_state.reglas_visibilidad
    )

# ==========================================================================================
# FIN PARTE 4/5
# (Sigue PARTE 5/5: Ajustes finales + validadores + recomendaciones de uso Survey123)
# ==========================================================================================
# ==========================================================================================
# PARTE 5/5 — AJUSTES FINALES + VALIDACIONES IMPORTANTES + RECOMENDACIONES DE USO
# ==========================================================================================

st.markdown("---")
st.subheader("✅ Chequeos rápidos (antes de exportar)")

def _check_consent_exists():
    return any(q.get("name") == "acepta_participar" for q in st.session_state.preguntas)

def _check_cascade_ready():
    # Verifica si hay al menos un cantón real y un distrito real en choices_ext_rows
    # (ignorando placeholders)
    if not st.session_state.choices_ext_rows:
        return False
    has_canton = any(r.get("list_name") == "list_canton" and r.get("name") not in (None, "__pick_canton__") for r in st.session_state.choices_ext_rows)
    has_distrito = any(r.get("list_name") == "list_distrito" and r.get("name") not in (None, "__pick_distrito__") for r in st.session_state.choices_ext_rows)
    return has_canton and has_distrito

def _check_names_unique():
    names = [q.get("name") for q in st.session_state.preguntas if q.get("name")]
    return len(names) == len(set(names))

c_ok = _check_consent_exists()
n_ok = _check_names_unique()
cas_ok = _check_cascade_ready()

col_a, col_b, col_c = st.columns(3)
with col_a:
    st.metric("Consentimiento (acepta_participar)", "OK" if c_ok else "FALTA")
    if not c_ok:
        st.warning("No existe `acepta_participar`. Sin esto, NO se puede filtrar por consentimiento.")
with col_b:
    st.metric("Names únicos", "OK" if n_ok else "DUPLICADOS")
    if not n_ok:
        st.warning("Hay 'name' duplicados. Edita preguntas hasta que cada 'name' sea único.")
with col_c:
    st.metric("Cascada Cantón→Distrito", "OK" if cas_ok else "INCOMPLETA")
    if not cas_ok:
        st.info("Agrega al menos 1 lote Cantón→Distrito en el catálogo para que la cascada funcione.")

st.markdown("---")
st.subheader("📌 Recomendaciones para Survey123 Connect")

st.markdown("""
1. **Exporta el XLSForm** con el botón **🧮 Construir XLSForm**.
2. En **Survey123 Connect**, crea/abre un proyecto → **New Survey** → **File** → selecciona el Excel.
3. Copia el logo a la carpeta `media/` del proyecto y asegúrate que el nombre coincida con `media::image`.
4. Prueba que:
   - **P2 Consentimiento** aparece después de la Intro.
   - Si marcás **No**, la encuesta muestra el mensaje de fin y **no** deja seguir.
   - Si marcás **Sí**, aparecen las páginas siguientes.
   - Cantón/Distrito cargan desde `choices` y Distrito filtra por Cantón.
5. Publica cuando todo esté OK.
""")

# ------------------------------------------------------------------------------------------
# Nota técnica del error que viste (para que quede trazabilidad)
# ------------------------------------------------------------------------------------------
with st.expander("🧠 Nota técnica: ¿por qué salía el error de conversión XLSForm?", expanded=False):
    st.markdown(f"""
El error:

> *There has been a problem trying to replace ${{acepta_participar}} ... There is no survey element...*

pasa cuando Survey123 intenta compilar un `relevant` que referencia **${{acepta_participar}}**
pero **esa pregunta todavía no existe** en el orden de la hoja `survey`.

✅ En esta versión se corrige porque:
- `acepta_participar` se escribe dentro de **P2 Consentimiento** ANTES de usarla.
- El “fin temprano” se maneja con un **grupo** con `relevant = ${{acepta_participar}}='no'`,
  mostrando un note y dejando que el flujo termine naturalmente.
- Todas las páginas siguientes tienen `relevant = ${{acepta_participar}}='si'` a nivel de grupo.
""")

# ==========================================================================================
# FIN DEL ARCHIVO — APP COMPLETA EN 5 PARTES
# ==========================================================================================
